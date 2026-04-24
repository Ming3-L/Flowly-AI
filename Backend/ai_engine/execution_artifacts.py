"""
从 ``WorkflowExecution.output_data`` 解析可导出 / 可预览的产物（文章、媒体 URL 等）。

画布形态：``{ "entry", "outputs": { node_id: { ... } }, "trace": [...] }``  
其它形态：尽力读取 ``response``、``text`` 等常见字段。
"""

from __future__ import annotations

import json
import re
from typing import Any
from urllib.parse import urlparse

_HTTP_URL = re.compile(r"^https?://", re.I)


def _is_http_url(s: str) -> bool:
    return bool(s and _HTTP_URL.match(s.strip()))


def _guess_kind_from_url(url: str) -> str | None:
    path = urlparse(url).path.lower()
    if any(path.endswith(ext) for ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp")):
        return "image"
    if any(path.endswith(ext) for ext in (".mp3", ".wav", ".ogg", ".m4a", ".aac", ".flac")):
        return "audio"
    if any(path.endswith(ext) for ext in (".mp4", ".webm", ".mov", ".mkv")):
        return "video"
    return None


def extract_primary_article_text(output_data: dict[str, Any]) -> str:
    """合并为可下载 / 预览的一篇正文（优先画布各节点 text，再兜底 JSON）。"""
    if not output_data:
        return ""

    outs = output_data.get("outputs")
    if isinstance(outs, dict) and outs:
        parts: list[str] = []
        for _nid, obj in outs.items():
            if not isinstance(obj, dict):
                continue
            t = obj.get("text")
            if isinstance(t, str) and t.strip():
                parts.append(t.strip())
        if parts:
            return "\n\n---\n\n".join(parts)

    for key in ("response", "content", "answer", "summary"):
        v = output_data.get(key)
        if isinstance(v, str) and v.strip():
            return v.strip()

    messages = output_data.get("messages")
    if isinstance(messages, list) and messages:
        last = messages[-1]
        if isinstance(last, dict):
            c = last.get("content")
            if isinstance(c, str) and c.strip():
                return c.strip()

    try:
        return json.dumps(output_data, ensure_ascii=False, indent=2)
    except Exception:
        return str(output_data)[:100_000]


def collect_media_artifacts(output_data: dict[str, Any]) -> dict[str, list[dict[str, Any]]]:
    """
    返回 ``{"images": [...], "audios": [...], "videos": [...]}``，
    每项 ``{"node_id", "url", "field"}``。
    """
    images: list[dict[str, Any]] = []
    audios: list[dict[str, Any]] = []
    videos: list[dict[str, Any]] = []

    outs = output_data.get("outputs") if isinstance(output_data, dict) else None
    if not isinstance(outs, dict):
        return {"images": images, "audios": audios, "videos": videos}

    def _push(lst: list[dict[str, Any]], node_id: str, url: str, field: str) -> None:
        url = url.strip()
        if any(x["url"] == url for x in lst):
            return
        lst.append({"node_id": node_id, "url": url, "field": field})

    for nid, obj in outs.items():
        if not isinstance(obj, dict):
            continue
        node_id = str(nid)

        for fld in ("image_url", "imageUrl"):
            raw = obj.get(fld)
            if isinstance(raw, str) and _is_http_url(raw):
                _push(images, node_id, raw, fld)

        au = obj.get("audio_url")
        if isinstance(au, str) and _is_http_url(au):
            _push(audios, node_id, au, "audio_url")

        vu = obj.get("video_url")
        if isinstance(vu, str) and _is_http_url(vu):
            _push(videos, node_id, vu, "video_url")

        gu = obj.get("url")
        if isinstance(gu, str) and _is_http_url(gu):
            kind = _guess_kind_from_url(gu)
            if kind == "image":
                _push(images, node_id, gu, "url")
            elif kind == "audio":
                _push(audios, node_id, gu, "url")
            elif kind == "video":
                _push(videos, node_id, gu, "url")

    return {"images": images, "audios": audios, "videos": videos}


def build_docx_bytes(title: str, body: str) -> bytes:
    from io import BytesIO

    from docx import Document  # type: ignore[import-untyped]
    from docx.shared import Pt  # type: ignore[import-untyped]

    doc = Document()
    doc.add_heading(title[:200] or "Flowly 导出", level=0)
    for block in body.split("\n\n"):
        para = doc.add_paragraph(block.replace("\r\n", "\n"))
        for run in para.runs:
            run.font.size = Pt(11)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
