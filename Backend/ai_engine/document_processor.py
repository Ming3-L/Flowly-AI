"""
Document Processor — Phase 8: RAG

Unified document processing pipeline supporting PDF, Word, TXT, HTML, Markdown, CSV.
Extracts text and metadata from uploaded files for chunking and embedding.
"""

from __future__ import annotations

import os
from pathlib import Path
from typing import Optional

from langchain_core.documents import Document  # pyright: ignore[reportMissingImports]


class DocumentProcessor:
    """
    Unified document processing for RAG knowledge bases.

    Supports: PDF, DOCX, TXT, HTML, Markdown, CSV
    """

    SUPPORTED_TYPES: set[str] = {"pdf", "docx", "txt", "html", "md", "csv"}

    def __init__(self, upload_dir: Optional[str] = None):
        self.upload_dir = upload_dir or os.getenv(
            "DOCUMENT_UPLOAD_DIR", "/data/uploads"
        )
        os.makedirs(self.upload_dir, exist_ok=True)

    def process(self, file_path: str) -> Document:
        """
        Process a file based on its extension and return a LangChain Document.

        Args:
            file_path: Absolute path to the file.

        Returns:
            A LangChain Document with page_content and metadata.

        Raises:
            ValueError: If the file type is not supported.
        """
        ext = Path(file_path).suffix.lstrip(".").lower()

        processor_map = {
            "pdf": self._process_pdf,
            "docx": self._process_docx,
            "txt": self._process_txt,
            "html": self._process_html,
            "md": self._process_markdown,
            "csv": self._process_csv,
        }

        if ext not in processor_map:
            raise ValueError(
                f"Unsupported file type: {ext}. "
                f"Supported: {', '.join(sorted(self.SUPPORTED_TYPES))}"
            )

        return processor_map[ext](file_path)

    def _process_pdf(self, file_path: str) -> Document:
        """Extract text and metadata from a PDF using PyMuPDF."""
        import fitz  # PyMuPDF  # pyright: ignore[reportMissingImports]

        doc = fitz.open(file_path)
        pages_text: list[str] = []

        for page_num, page in enumerate(doc):
            text = page.get_text("text") or ""
            if text.strip():
                pages_text.append(text.strip())

        metadata = {
            "source": file_path,
            "total_pages": len(doc),
            "file_type": "pdf",
        }

        # Extract PDF metadata
        pdf_meta = doc.metadata
        if pdf_meta:
            metadata["title"] = pdf_meta.get("title", "")
            metadata["author"] = pdf_meta.get("author", "")
            metadata["subject"] = pdf_meta.get("subject", "")

        doc.close()

        return Document(
            page_content="\n\n".join(pages_text),
            metadata=metadata,
        )

    def _process_docx(self, file_path: str) -> Document:
        """Extract text from a Word DOCX file using python-docx."""
        try:
            import docx  # pyright: ignore[reportMissingImports]
        except ImportError:
            return Document(
                page_content="[python-docx not installed — install with: pip install python-docx]",
                metadata={"source": file_path, "file_type": "docx", "error": "missing_dependency"},
            )

        doc_obj = docx.Document(file_path)
        paragraphs: list[str] = []

        for para in doc_obj.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Extract core properties
        core_props = doc_obj.core_properties
        metadata = {
            "source": file_path,
            "file_type": "docx",
            "title": core_props.title or "",
            "author": core_props.author or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
            "paragraph_count": len(paragraphs),
        }

        return Document(
            page_content="\n\n".join(paragraphs),
            metadata=metadata,
        )

    def _process_txt(self, file_path: str) -> Document:
        """Extract plain text from a TXT file."""
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            content = f.read()

        stat = os.stat(file_path)
        return Document(
            page_content=content,
            metadata={
                "source": file_path,
                "file_type": "txt",
                "file_size": stat.st_size,
                "line_count": content.count("\n") + 1,
            },
        )

    def _process_html(self, file_path: str) -> Document:
        """Extract visible text from an HTML file."""
        try:
            from bs4 import BeautifulSoup
        except ImportError:
            return self._process_txt(file_path)

        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            html_content = f.read()

        soup = BeautifulSoup(html_content, "html.parser")

        # Remove script and style elements
        for tag in soup(["script", "style", "noscript"]):
            tag.decompose()

        text = soup.get_text(separator="\n", strip=True)
        # Collapse multiple blank lines
        import re
        text = re.sub(r"\n{3,}", "\n\n", text)

        return Document(
            page_content=text,
            metadata={
                "source": file_path,
                "file_type": "html",
                "title": soup.title.string if soup.title else "",
            },
        )

    def _process_markdown(self, file_path: str) -> Document:
        """Process a Markdown file as plain text (metadata extracted from headers)."""
        content = self._process_txt(file_path)
        lines = content.page_content.split("\n")

        metadata = {**content.metadata}

        # Extract title from first H1
        for line in lines:
            if line.startswith("# "):
                metadata["title"] = line[2:].strip()
                break

        metadata["file_type"] = "md"
        return Document(page_content=content.page_content, metadata=metadata)

    def _process_csv(self, file_path: str) -> Document:
        """Convert CSV rows to a plain-text document."""
        import csv

        rows: list[str] = []
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            reader = csv.reader(f)
            for row in reader:
                rows.append(", ".join(row))

        stat = os.stat(file_path)
        return Document(
            page_content="\n".join(rows),
            metadata={
                "source": file_path,
                "file_type": "csv",
                "file_size": stat.st_size,
                "row_count": len(rows),
            },
        )

    def process_batch(self, file_paths: list[str]) -> list[Document]:
        """
        Process multiple files and return a list of Documents.

        Args:
            file_paths: List of absolute file paths.

        Returns:
            List of Documents, one per file. Failed files are skipped with an error doc.
        """
        results: list[Document] = []
        for path in file_paths:
            try:
                doc = self.process(path)
                results.append(doc)
            except Exception as exc:
                results.append(
                    Document(
                        page_content=f"[Processing failed for {path}: {exc}]",
                        metadata={"source": path, "file_type": "unknown", "error": str(exc)},
                    )
                )
        return results
